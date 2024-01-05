from django.http import HttpResponse
from docx import Document
from rest_framework import status
from rest_framework.parsers import FormParser, MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView

from .for_text import text_convert_to_latin, text_convert_to_cyrillic
from .serializer import DocxSerializer, TxtFileSerializer
from conventer.for_docx import docx_convert_to_latin, docx_conventer_to_cyrillic


class DocxLatinAPIView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    serializer_class = DocxSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data)
        if serializer.is_valid():
            serializer.save()
            yolak = serializer.data['docx_file']
            yolak = str(yolak)
            print(yolak)
            yolak = yolak.replace("/media/docx_files/", "")
            text = ""
            document = Document(f"C://Users/user/Desktop/ConventerApi/mediafiles/docx_files/{yolak}")
            for i in document.paragraphs:
                text += i.text
            response = text_convert_to_latin(text)
            return Response({
                "result": response
            })

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


class DocxCyrillicAPIView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    serializer_class = DocxSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data)
        if serializer.is_valid():
            serializer.save()
            yolak = serializer.data['docx_file']
            yolak = str(yolak)
            yolak = yolak.replace("/media/docx_files/", "")
            text = ""
            document = Document(f"C://Users/user/Desktop/ConventerApi/mediafiles/docx_files/{yolak}")
            for i in document.paragraphs:
                text += i.text
            response = text_convert_to_cyrillic(text)
            return Response({
                "result": response
            })

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


class TxtCyrillicAPIView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    serializer_class = TxtFileSerializer

    def post(self, request, *args, **kwargs):
        text = ""
        serializer = self.serializer_class(data=request.data)
        if serializer.is_valid():
            serializer.save()
            yolak = serializer.data['txt']
            yolak = str(yolak)
            yolak = yolak.replace("/media/txt_files/", "")
            f = open(f"C://Users/user/Desktop/ConventerApi/mediafiles/txt_files/{yolak}", "r", encoding="utf-8")
            text = f.read()
            response = text_convert_to_cyrillic(text)
            return Response({
                "result": response
            })

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


class TxtLatinAPIView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    serializer_class = TxtFileSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data)
        if serializer.is_valid():
            serializer.save()
            yolak = serializer.data['txt']
            yolak = str(yolak)
            print(yolak)
            yolak = yolak.replace("/media/txt_files/", "")
            f = open(f"C://Users/user/Desktop/ConventerApi/mediafiles/txt_files/{yolak}", "r", encoding="utf-8")
            text = f.read()
            response = text_convert_to_latin(text)
            return Response({
                "result": response
            })

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


class TextView(APIView):
    def post(self, request):
        context = request.data['context']
        pattern = request.data['pattern']
        if pattern == 'latin':
            response = text_convert_to_latin(context)
        elif pattern == 'cyrillic':
            response = text_convert_to_cyrillic(context)
        else:
            return Response({
                "Message": "Wrong Input",
                "You can choose only": "cyrillic or latin"
            })

        return Response({
            "result": f"{response}"
        })


def migration(request):
    import os
    os.system('python3 manage.py makemigrations')
    os.system('python3 manage.py migrate --no-input')
    return HttpResponse('Migration Done')
